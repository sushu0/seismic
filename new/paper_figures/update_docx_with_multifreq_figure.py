from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create a new docx with the updated multi-frequency thin-layer figure attached.")
    parser.add_argument("--source-docx", required=True)
    parser.add_argument("--figure-png", required=True)
    parser.add_argument("--output-docx", required=True)
    return parser.parse_args()


def style_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)


def main() -> None:
    args = parse_args()
    source_docx = Path(args.source_docx)
    figure_png = Path(args.figure_png)
    output_docx = Path(args.output_docx)

    if not source_docx.exists():
        raise FileNotFoundError(f"Source docx not found: {source_docx}")
    if not figure_png.exists():
        raise FileNotFoundError(f"Figure png not found: {figure_png}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, output_docx)

    document = Document(str(output_docx))
    document.add_section(WD_SECTION.NEW_PAGE)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("图件更新说明")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    note = document.add_paragraph(
        "原始 Word 文件中未检测到可稳定定位的旧图锚点或内嵌媒体关系。"
        "为避免误改正文结构，这里附入新的“浅部薄层自建模型不同频率反演结果图”及替换说明，"
        "建议在排版稿中用下图替换原对应旧图，并保持原图号与正文引用不变。"
    )
    note.paragraph_format.space_after = Pt(6)

    picture_para = document.add_paragraph()
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.add_run().add_picture(str(figure_png), width=Cm(16.5))

    caption = document.add_paragraph(
        "建议替换图题：浅部薄层自建模型不同频率反演结果对比。"
        "自左上至右下分别为参考真值、20 Hz 结果、30 Hz 结果和 40 Hz 结果。"
    )
    style_caption(caption)

    detail = document.add_paragraph(
        "替换建议：将此页插图替换到原“浅部薄层自建模型不同频率反演结果图”所在位置；"
        "若原图题已存在，请保留原图号，仅更新图片内容与必要的图题措辞。"
    )
    detail.paragraph_format.space_before = Pt(3)

    document.save(str(output_docx))
    print(f"Saved docx: {output_docx}")


if __name__ == "__main__":
    main()
